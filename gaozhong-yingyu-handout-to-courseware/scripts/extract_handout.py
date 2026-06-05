# -*- coding: utf-8 -*-
"""
extract_handout.py — 高中英语讲义 docx → content.json（A类内容，100%讲义直出）

自动识别讲次类型：
  continuation 读后续写(主观题)：含「固化情节」标题 → part1_plots + part2/part3(语篇+续写提示+情节描写)
  reading      阅读理解(客观题)：含「阅读X篇」标题 → part1_zhuzhi/part1_xuanxiang + part2[篇...]
同时抽取讲义内嵌图(Knowledge Map) → assets/knowledge_map.png

用法: python3 extract_handout.py <讲义.docx> -o <outdir>
"""
import argparse, json, re, shutil, zipfile
from pathlib import Path
import docx
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table


def load_items(src):
    doc = docx.Document(src)
    items = []
    for ch in doc.element.body.iterchildren():
        if isinstance(ch, CT_P):
            p = Paragraph(ch, doc)
            items.append(("p", p.style.name if p.style else "", p.text.strip()))
        elif isinstance(ch, CT_Tbl):
            t = Table(ch, doc)
            rows = [[c.paragraphs[0].text.strip() if c.paragraphs else "" for c in r.cells]
                    for r in t.rows]
            items.append(("t", "", rows))
    return items


def extract_km(src, outdir):
    z = zipfile.ZipFile(src)
    imgs = [n for n in z.namelist() if n.startswith("word/media/")
            and "header" not in n and "footer" not in n
            and n.lower().endswith((".png", ".jpg", ".jpeg"))]
    if not imgs:
        return None
    out = Path(outdir) / "assets" / "knowledge_map.png"
    out.parent.mkdir(parents=True, exist_ok=True)
    with z.open(imgs[0]) as f, open(out, "wb") as o:
        shutil.copyfileobj(f, o)
    return str(out)


def base_meta(items):
    title_line = next((t for k, st, t in items if k == "p" and t), "")
    m = re.match(r'^第?([一二三四五六七八九十0-9]+)讲\s*(.*)$', title_line.replace(" ", " ").strip())
    num_map = {"一": "1", "二": "2", "三": "3", "四": "4", "五": "5", "六": "6",
               "七": "7", "八": "8", "九": "9", "十": "10", "一十": "10"}
    lecture_no, title = "第X讲", title_line
    if m:
        n = m.group(1)
        n = num_map.get(n, re.sub(r'[一二三四五六七八九十]', lambda x: num_map.get(x.group(), ""), n)) or n
        lecture_no = f"第{int(n):02d}讲" if n.isdigit() else f"第{n}讲"
        title = m.group(2).strip()
    C = {"lecture_no": lecture_no,
         "title": "\n".join([title[:len(title) // 2], title[len(title) // 2:]]) if len(title) >= 6 else title,
         "grade": "高中英语", "term": "", "teacher": "主讲老师：",
         "end": {"big": "本节课结束", "small": "下节课我们\n再见啦～"}}
    for it in items:                       # Preview 表
        if it[0] == "t" and it[2] and it[2][0] and it[2][0][0] == "Parts":
            C["preview"] = {"header": it[2][0], "rows": it[2][1:]}
            break
    for k, it in enumerate(items):         # Leading-in
        if it[0] == "p" and it[2] == "Leading-in":
            for j in range(k + 1, len(items)):
                if items[j][0] == "p" and items[j][2]:
                    C["leading_in"] = items[j][2]; break
            break
    return C


def heads(items, level):
    return [(i, it[2]) for i, it in enumerate(items)
            if it[0] == "p" and it[1].startswith(f"Heading {level}")]


def seg(items, start, end):
    return items[start + 1:end]


# ───────── reading（客观题·阅读）─────────
def is_stem(t): return bool(re.match(r'^(无\d+\s*)?(Q\d+\s*[:：]|\(\d+\)|（\d+）)', t))
def is_opt(t): return bool(re.match(r'^[A-D][．.、]', t))


def parse_source_line(txt):
    txt = re.sub(r'^无\d+\s*', '', txt)
    stars = ''.join(re.findall(r'[⭐★]', txt[:12]))
    txt = re.sub(r'^[⭐★️\s]+', '', txt)
    m = re.search(r'【[^】]*】', txt)
    src, rest = "", txt
    if m:
        src = m.group(0); rest = (txt[:m.start()] + txt[m.end():]).strip()
    return stars, src, rest


def collect_reading(paras):
    level = source = ""; passage = []; i = 0
    while i < len(paras) and not paras[i]:
        i += 1
    if i < len(paras):
        level, source, rest = parse_source_line(paras[i])
        if rest:
            passage.append(rest)
        i += 1
    while i < len(paras) and not is_stem(paras[i]):
        t = paras[i]
        if t and not t.startswith("【") and not is_opt(t):
            passage.append(t)
        i += 1
    questions, answers = [], []
    cur = None; atag = aans = aana = None; mode = None
    while i < len(paras):
        t = paras[i]; i += 1
        if not t:
            continue
        if is_stem(t):
            if cur: questions.append(cur)
            cur = {"stem": re.sub(r'^无\d+\s*', '', t), "options": []}; mode = "q"
        elif is_opt(t) and cur and mode == "q":
            cur["options"].append(t)
        elif t.startswith("【知识标签】"):
            if cur: questions.append(cur); cur = None
            mode = "a"
            if atag is not None:
                answers.append((atag, aans, aana))
            atag = t[6:].strip(); aans = aana = ""
        elif t.startswith("【答案】"):
            aans = t[4:].strip()
        elif t.startswith("【解析】"):
            aana = t[4:].strip()
        elif t.startswith(("【小题答案】", "【主旨辅助】", "_")):
            continue
        elif mode == "a" and aana is not None:
            aana += t
    if cur: questions.append(cur)
    if atag is not None:
        answers.append((atag, aans, aana))
    real = [a for a in answers if a[1]]
    for q, (tg, an, ana) in zip(questions, real):
        q["tag"], q["answer"], q["analysis"] = tg, an, ana
    return source, level, passage, questions


def parse_reading(items, C):
    C["lecture_type"] = "reading"
    h3 = heads(items, 3)
    names = [n for _, n in h3]
    def section(name):
        for k, (i, n) in enumerate(h3):
            if n == name:
                end = h3[k + 1][0] if k + 1 < len(h3) else len(items)
                return items[i + 1:end]
        return []
    def paras_of(sec): return [it[2] for it in sec if it[0] == "p"]
    def tables_of(sec): return [it[2] for it in sec if it[0] == "t"]

    # PART1：主旨辅助 / 选项辅助（若存在）
    if "主旨辅助" in names:
        sec = section("主旨辅助")
        src, lv, psg, qs = collect_reading(paras_of(sec))
        tbls = tables_of(sec)
        C["part1_zhuzhi"] = {"title": "主旨辅助", "source": src, "passage": psg,
                             "method_table": tbls[-1] if tbls else [], "questions": qs}
    if "选项辅助" in names:
        tbls = tables_of(section("选项辅助"))
        C["part1_xuanxiang"] = {"title": "选项辅助", "table": tbls[-1] if tbls else []}
    # PART2：阅读X篇
    C["part2"] = []
    for _, n in h3:
        if re.match(r'^阅读.+篇$', n):
            src, lv, psg, qs = collect_reading(paras_of(section(n)))
            C["part2"].append({"name": n, "level": lv, "source": src,
                               "passage": psg, "questions": qs})
    C["parts"] = [
        {"part_label": "PART 1", "title": "辅助方法", "subtitle": " · ".join(
            [t for t in ["主旨辅助" if "主旨辅助" in names else "",
                         "选项辅助" if "选项辅助" in names else ""] if t])},
        {"part_label": "PART 2", "title": "篇章训练",
         "subtitle": " · ".join(p["name"] for p in C["part2"])},
    ]
    return C


# ───────── continuation（主观题·读后续写）─────────
def parse_continuation(items, C):
    C["lecture_type"] = "continuation"
    h3 = heads(items, 3)

    # PART1 固化情节N：基础/升级 中英
    plots = []
    for k, (i, name) in enumerate(h3):
        m = re.match(r'^固化情节[一二三四五六七八九十]+[：:](.+)$', name)
        if not m:
            continue
        end = h3[k + 1][0] if k + 1 < len(h3) else len(items)
        paras = [it[2] for it in items[i + 1:end] if it[0] == "p" and it[2]]
        plot = {"index": len(plots) + 1, "name": m.group(1).strip(),
                "basic": {}, "upgrade": {}}
        slot = None
        for t in paras:
            if t.startswith("1基础") or "基础版本" in t[:6]:
                slot = "basic"
            elif t.startswith("2升级") or "升级版本" in t[:6]:
                slot = "upgrade"
            elif t.startswith("【答案】中文") and slot:
                plot[slot]["zh"] = t.split("：", 1)[-1].split(":", 1)[-1].strip()
            elif t.startswith("英文") and slot:
                plot[slot]["en"] = t.split("：", 1)[-1].split(":", 1)[-1].strip()
            elif t.startswith("【知识标签】") and slot:
                plot[slot]["tag"] = t
        if plot["basic"].get("zh"):
            plots.append(plot)
    C["part1_plots"] = plots

    # PART2/3：篇章（H1 = 篇章实战 / 篇章复用）
    h1 = heads(items, 1)
    def part_block(h1name):
        for k, (i, n) in enumerate(h1):
            if n == h1name:
                end = h1[k + 1][0] if k + 1 < len(h1) else len(items)
                return items[i + 1:end]
        return []

    def parse_part(block):
        paras = [it[2] for it in block if it[0] == "p"]
        part = {"source": "", "passage_en": [], "prompts": {}, "details": []}
        i = 0
        while i < len(paras):
            t = paras[i]; i += 1
            if not t:
                continue
            if not part["source"] and re.match(r'^【\d{4}', t):
                part["source"] = re.search(r'【[^】]*】', t).group(0)
                rest = re.sub(r'【[^】]*】', '', t).strip()
                if rest:
                    part["passage_en"].append(rest)
            elif re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', t):
                part["passage_en"].append(t)
            elif t.startswith("Para 1:"):
                part["prompts"]["para1"] = t.rstrip("_ ").strip()
            elif t.startswith("Para 2:"):
                part["prompts"]["para2"] = t.rstrip("_ ").strip()
            else:
                m = re.match(r'^(\d)【(.+?)】中文[：:](.+)$', t)
                if m:
                    part["details"].append({"index": int(m.group(1)),
                                            "group": "首段" if int(m.group(1)) <= 4 else "次段",
                                            "label": m.group(2), "zh": m.group(3).strip()})
                elif t.startswith("【答案】") and part["details"] and \
                        "en" not in part["details"][-1]:
                    part["details"][-1]["en"] = t[4:].strip()
                elif t.startswith("【知识标签】") and part["details"]:
                    part["details"][-1].setdefault("tag", t)
        return part

    for h1name, key in [("篇章实战", "part2"), ("篇章复用", "part3")]:
        blk = part_block(h1name)
        if blk:
            C[key] = parse_part(blk)
    C["parts"] = [
        {"part_label": "PART 1", "title": "固化情节", "subtitle": "固化情节内容提炼"},
        {"part_label": "PART 2", "title": "篇章实战", "subtitle": "固化情节篇章实战"},
        {"part_label": "PART 3", "title": "篇章复用", "subtitle": "内容改写迁移复用"},
    ]
    return C


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("src")
    ap.add_argument("-o", "--outdir", required=True)
    a = ap.parse_args()
    outdir = Path(a.outdir); outdir.mkdir(parents=True, exist_ok=True)
    items = load_items(a.src)
    C = base_meta(items)
    all_h3 = " ".join(n for _, n in heads(items, 3))
    if "固化情节" in all_h3:
        C = parse_continuation(items, C)
    else:
        C = parse_reading(items, C)
    km = extract_km(a.src, outdir)
    (outdir / "content.json").write_text(json.dumps(C, ensure_ascii=False, indent=2),
                                         encoding="utf-8")
    print(f"✅ content.json（type={C['lecture_type']}）→ {outdir/'content.json'}")
    if km:
        print(f"✅ Knowledge Map 图 → {km}")
    if C["lecture_type"] == "continuation":
        print(f"   固化情节×{len(C.get('part1_plots', []))}；"
              f"part2 details×{len(C.get('part2', {}).get('details', []))}；"
              f"part3 details×{len(C.get('part3', {}).get('details', []))}")
    else:
        print(f"   主旨辅助Q×{len(C.get('part1_zhuzhi', {}).get('questions', []))}；"
              f"篇章×{len(C.get('part2', []))}")


if __name__ == "__main__":
    main()
